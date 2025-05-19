from docx import Document
from docx.shared import Pt
from io import BytesIO

def create_document_bytes(content):
    doc = Document()
    title = content[0]['text']
    h0 = doc.add_heading(title, level=0)
    h0.runs[0].font.size = Pt(14)
    
    for item in content[1:]:
        if item['type'] == 'header':
            h = doc.add_heading(level=item['level'])
            run = h.add_run(item['text'])
            run.bold = True
            run.font.size = Pt(14)
        elif item['type'] == 'paragraph':
            p = doc.add_paragraph()
            if 'parts' in item:
                for part in item['parts']:
                    run = p.add_run(part['text'])
                    if part['bold']:
                        run.bold = True
                    run.font.size = Pt(12)
            else:
                run = p.add_run(item['text'])
                run.font.size = Pt(12)
        elif item['type'] == 'list':
            for li in item['items']:
                p = doc.add_paragraph(style='List Bullet')
                if 'parts' in li:
                    for part in li['parts']:
                        run = p.add_run(part['text'])
                        if part['bold']:
                            run.bold = True
                        run.font.size = Pt(12)
                else:
                    run = p.add_run(li)
                    run.font.size = Pt(12)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
